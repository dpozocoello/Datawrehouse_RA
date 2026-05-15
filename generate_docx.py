import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: El archivo {md_path} no existe.")
        return

    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            content = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(content, style='List Number')
        elif line.startswith('|'):
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
        elif line.startswith('```'):
            continue
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    try:
        doc.save(docx_path)
        print(f"Archivo generado exitosamente en: {docx_path}")
    except Exception as e:
        # Si falla por estar abierto, intentar con otro nombre
        new_path = docx_path.replace(".docx", "_v2.docx")
        try:
            doc.save(new_path)
            print(f"Original bloqueado. Generado como: {new_path}")
        except:
            print(f"Fallo crítico al guardar: {e}")

if __name__ == "__main__":
    files = [
        (r"f:\Datawrehouse_RA\documentación\ETL_Arquitectura_Estratégica_Detallada.md", 
         r"f:\Datawrehouse_RA\documentación\ETL_Arquitectura_Estratégica_Detallada.docx"),
        (r"f:\Datawrehouse_RA\documentación\Manual_Uso_DWH_Analitica.md", 
         r"f:\Datawrehouse_RA\documentación\Manual_Uso_DWH_Analitica.docx")
    ]
    
    for md, docx in files:
        markdown_to_docx(md, docx)
