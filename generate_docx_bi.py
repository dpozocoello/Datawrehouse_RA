import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: El archivo {md_path} no existe.")
        return

    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Saltos de línea vacíos
        if not line:
            doc.add_paragraph()
            continue

        # Encabezados
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        
        # Listas
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            content = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(content, style='List Number')
        
        # Tablas
        elif line.startswith('|'):
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New'
        
        # Bloques de código
        elif line.startswith('```'):
            continue
            
        # Párrafos normales
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    try:
        doc.save(docx_path)
        print(f"Archivo generado exitosamente en: {docx_path}")
    except Exception as e:
        print(f"Error al guardar {docx_path}: {e}")

if __name__ == "__main__":
    md_file = r"f:\Datawrehouse_RA\documentación\Manual_Uso_DWH_Analitica.md"
    docx_file = r"f:\Datawrehouse_RA\documentación\Manual_Uso_DWH_Analitica.docx"
    markdown_to_docx(md_file, docx_file)
