import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_styled_report(md_file_path, output_docx_path):
    """
    Convierte el borrador maestro a un documento Word (.docx) con estilo corporativo.
    """
    doc = docx.Document()

    # --- Configuración Estilo General ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Portada Simple ---
    title = doc.add_heading('Informe de Gestión DWH v1.6', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Regularización Ambiental – Ecosistema de Datos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_paragraph('\n' * 5)
    
    meta_info = doc.add_paragraph()
    meta_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_info.add_run('Fecha: Marzo 2026\n')
    meta_info.add_run('Autor: Antigravity AI – Senior Data Architect\n')
    meta_info.add_run('Versión: 1.6\n')

    doc.add_page_break()

    # --- Procesamiento de Contenido (Basado en el MD) ---
    if not os.path.exists(md_file_path):
        print(f"Error: {md_file_path} no encontrado.")
        return

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Saltar título ya puesto en portada
        if line.startswith('# Informe de Gestión'):
            continue
        
        # Headers
        if line.startswith('## '):
            if in_table: # Cerrar tabla si existía
                create_table(doc, current_table)
                current_table = []
                in_table = False
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        
        # Tablas
        elif line.startswith('|'):
            in_table = True
            if '---' in line: # Saltar separador de MD
                continue
            current_table.append([cell.strip() for cell in line.split('|') if cell.strip()])
        
        # Párrafos y Listas
        elif line:
            if in_table:
                create_table(doc, current_table)
                current_table = []
                in_table = False
            
            p = doc.add_paragraph()
            # Negritas en MD (**texto**)
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
            
            # Listas de MD (- )
            if line.startswith('- '):
                p.style = 'List Bullet'
                p.runs[0].text = p.runs[0].text.replace('- ', '', 1)

    # El último bloque si fue tabla
    if in_table and current_table:
        create_table(doc, current_table)

    # --- Firma ---
    doc.add_paragraph('\n' * 2)
    footer = doc.add_paragraph('__________________________\n')
    footer.add_run('Antigravity AI – Arquitectura de Datos\n')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_docx_path)
    print(f"Documento generado exitosamente en: {output_docx_path}")

def create_table(doc, data):
    if not data: return
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
    doc.add_paragraph() # Espacio después de tabla

if __name__ == "__main__":
    md_path = r"D:\Datawrehouse_RA\documentación\informe_de_gestion_v1_6.md"
    docx_path = r"D:\Datawrehouse_RA\documentación\Informe de Gestión DWH v1.6.docx"
    create_styled_report(md_path, docx_path)
