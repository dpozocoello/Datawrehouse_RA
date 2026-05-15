import os
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_styled_architecture_doc(md_file_path, output_docx_path):
    """
    Convierte el borrador maestro de arquitectura a un documento Word (.docx) profesional.
    """
    doc = docx.Document()

    # --- Configuración Estilo General ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Portada Profesional ---
    # Título Principal
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = t.add_run('ARQUITECTURA TÉCNICA Y ESTRATEGIA DE DATOS')
    run_t.font.size = Pt(24)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph('\n')

    # Subtítulo
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = st.add_run('Data Warehouse Regularización Ambiental (v1.6)')
    run_st.font.size = Pt(16)
    run_st.font.italic = True

    doc.add_paragraph('\n' * 8)

    # Información de Autoría
    meta_info = doc.add_paragraph()
    meta_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_info.add_run('Fecha: Marzo 2026\n').bold = True
    meta_info.add_run('Documento: Especificación de Ingeniería DWH\n')
    meta_info.add_run('Estado: Desarrollo (DEV) - Listo para Promoción\n')
    meta_info.add_run('Arquitecto: Antigravity AI – Senior Data Architect\n')

    doc.add_page_break()

    # --- Tabla de Contenidos (Simulada para Docx) ---
    doc.add_heading('Tabla de Contenidos', level=1)
    doc.add_paragraph('1. Cronología Evolutiva del Desarrollo\n'
                      '2. Fase de Análisis: Ecosistema Multiorigen\n'
                      '3. Fase de Diseño: Arquitectura de Datos\n'
                      '4. Fase de Preparación: Infraestructura y Entorno\n'
                      '5. Fase de Desarrollo: Ingeniería de Datos\n'
                      '6. Flujo de Promoción Ambiental (Deployment Plan)\n'
                      '7. Conclusión de Arquitectura')
    doc.add_page_break()

    # --- Procesamiento de Contenido (Basado en el MD) ---
    if not os.path.exists(md_file_path):
        print(f"Error: {md_file_path} no encontrado.")
        return

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = []
    in_table = False

    for line in lines:
        line = line.strip()
        
        # Saltar títulos de portada
        if line.startswith('# Arquitectura Técnica'):
            continue
        
        # Headers
        if line.startswith('## '):
            if in_table:
                create_table(doc, current_table)
                current_table = []
                in_table = False
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            if in_table:
                create_table(doc, current_table)
                current_table = []
                in_table = False
            doc.add_heading(line.replace('### ', ''), level=2)
        
        # Tablas
        elif line.startswith('|'):
            in_table = True
            if '---' in line:
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                current_table.append(cells)
        
        # Párrafos y Listas
        elif line:
            if in_table:
                create_table(doc, current_table)
                current_table = []
                in_table = False
            
            p = doc.add_paragraph()
            # Negritas e Itálicas simples
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
            
            if line.startswith('- '):
                p.style = 'List Bullet'
                p.runs[0].text = p.runs[0].text.replace('- ', '', 1)

    # Cierre de última tabla
    if in_table and current_table:
        create_table(doc, current_table)

    # --- Firma Final ---
    doc.add_paragraph('\n' * 3)
    footer = doc.add_paragraph('__________________________\n')
    footer.add_run('Antigravity AI – Senior Data Architect\n')
    footer.add_run('Unidad de Inteligencia de Datos').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_docx_path)
    print(f"Master Document generado: {output_docx_path}")

def create_table(doc, data):
    if not data: return
    try:
        # Validar consistencia de columnas
        cols = len(data[0])
        table = doc.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'
        
        for row_cells in data:
            row = table.add_row()
            for i, cell_text in enumerate(row_cells):
                if i < cols:
                    row.cells[i].text = cell_text
        doc.add_paragraph()
    except Exception as e:
        print(f"Error creando tabla: {e}")

if __name__ == "__main__":
    md_path = r"D:\Datawrehouse_RA\documentación\DWH_Documentacion_Tecnica_Total_v1_6.md"
    docx_path = r"D:\Datawrehouse_RA\documentación\Manual Arquitectura y Bitácora Total DWH RA.docx"
    create_styled_architecture_doc(md_path, docx_path)
